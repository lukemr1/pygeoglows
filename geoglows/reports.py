import datetime
import os
import io
import requests
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx2pdf import convert

import pandas as pd

from .data import *
from ._plots import plots

today = datetime.date.today()

# ==========================================
# INTERNAL HELPER FUNCTIONS
# ==========================================

def _add_cover_page(doc, report_type):
    """
    Adds a professional cover page with the GEOGloWS logo, title, and reference links.
    """
    # --- 1. GEOGloWS Logo ---
    logo_url = "https://training.geoglows.org/static/images/NewGEOGLOWSLOGO.png"
    try:
        response = requests.get(logo_url, timeout=5)
        if response.status_code == 200:
            logo_stream = io.BytesIO(response.content)
            logo_p = doc.add_paragraph()
            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_p.add_run()
            logo_run.add_picture(logo_stream, width=Inches(3))
            doc.add_paragraph("\n")
    except Exception:
        pass

    # --- 2. Title Section ---
    title_p = doc.add_paragraph("\n" * 6)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Hydrology Report")
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark Blue

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(f"{report_type}\nGenerated: {today.strftime('%B %d, %Y')}")
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True

    # --- 3. Spacer to push references to bottom ---
    doc.add_paragraph("\n" * 9)

    # --- 4. References & Sources ---
    doc.add_heading("Data Sources & References", level=2)

    sources = [
        ("Primary Forecast Engine", "RFS-Hydroviewer API https://hydroviewer.geoglows.org"),
        ("Historical Data", "AWS Public Dataset:\n s3://geoglows-v2-retrospective"),
        ("Training & Documentation", "https://training.geoglows.org"),
        ("More Information", "https://www.geoglows.org")
    ]

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light List Accent 1'

    for source, link in sources:
        row_cells = table.add_row().cells
        row_cells[0].text = source
        row_cells[1].text = link

    disclaimer = doc.add_paragraph(
        "\nThis report was generated automatically. Verify all results with local observation data.")
    disclaimer.style = "Caption"

    doc.add_page_break()


def _save_plots_to_file(figures_or_streams, filename, report_type, output_format='docx', input_is_bytes=False):
    """
    Helper function that accepts plotly figures (or byte streams) and outputs a DOCX or PDF report.
    """
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    _add_cover_page(doc, report_type)

    for i, item in enumerate(figures_or_streams):
        # Optimization: Handle pre-rendered bytes or raw figures
        if input_is_bytes:
            img_stream = item
        else:
            # Slow path: render figure here
            img_bytes = item.to_image(format="png", width=1000, height=600, scale=2)
            img_stream = io.BytesIO(img_bytes)

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(img_stream, width=Inches(6.0))

        caption = doc.add_paragraph(f"Figure {i + 1}")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n")

        doc.add_heading(f'Comments for Figure {i + 1}', level=3)
        doc.add_paragraph("Notes:")
        doc.add_paragraph("\n" * 3)
        doc.add_paragraph("_" * 50)
        doc.add_paragraph("\n")

        doc.add_page_break()

    doc.add_heading('Overall Comments', level=1)
    doc.add_paragraph("Please add any additional notes regarding this report below:")
    doc.add_paragraph("\n" * 10)

    downloads_path = str(Path.home() / "Downloads")

    if not filename.endswith('.docx'):
        filename = filename + '.docx'

    docx_path = os.path.join(downloads_path, filename)
    doc.save(docx_path)
    print(f"DOCX Saved: {docx_path}")

    if output_format.lower() == 'pdf':
        print("Converting to PDF (this may take a moment)...")
        pdf_path = docx_path.replace('.docx', '.pdf')
        try:
            convert(docx_path, pdf_path)
            os.remove(docx_path)
            abs_path = os.path.abspath(pdf_path)
        except Exception as e:
            print(f"PDF conversion failed: {e}. Saving as DOCX instead.")
            abs_path = os.path.abspath(docx_path)
        return abs_path
    else:
        abs_path = os.path.abspath(docx_path)
        return abs_path


def _add_return_period_table(doc, rp_df, ensemble_df, forecast_df):
    """
    Add return periods table with probabilities based on daily maximum flows.

    This matches the JavaScript approach: for each day, we find the maximum flow
    across all timesteps within that day, then calculate the probability that
    this daily maximum exceeds each return period threshold.
    """
    doc.add_heading('Return Period Thresholds', level=2)

    # Step 1: Convert ensemble_df index to datetime and get unique dates
    ensemble_df.index = pd.to_datetime(ensemble_df.index)

    # Get unique dates (up to 15 days)
    unique_dates = ensemble_df.index.normalize().unique()[:15]

    # Step 2: Calculate daily maximums for each ensemble member
    # ensemble_df has columns for each ensemble member (e.g., '0', '1', '2', ...)
    # We'll create a new dataframe where each row is a day and each column is an ensemble member's daily max

    daily_max_data = []
    forecast_dates = []

    for unique_date in unique_dates:
        # Get all timesteps for this date
        mask = ensemble_df.index.normalize() == unique_date
        day_data = ensemble_df[mask]

        if len(day_data) > 0:
            # For each ensemble member (column), get the maximum value for this day
            daily_max = day_data.max(axis=0)  # max across all timesteps for each member
            daily_max_data.append(daily_max)
            forecast_dates.append(unique_date)

    # Create a DataFrame where each row is a day, each column is an ensemble member
    daily_max_df = pd.DataFrame(daily_max_data, index=forecast_dates)

    num_dates = len(forecast_dates)

    # Step 3: Create the table structure
    table = doc.add_table(rows=len(rp_df) + 1, cols=num_dates + 1)
    table.style = 'Light Grid Accent 1'

    # Prevent row splitting across pages
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))

    # Set column widths
    table.columns[0].width = Inches(1.3)
    for i in range(1, num_dates + 1):
        table.columns[i].width = Inches(0.65)

    # Step 4: Create header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Return Periods'
    hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hdr_cells[0].paragraphs[0].runs:
        run.font.size = Pt(7)
        run.font.bold = True

    # Add date headers
    for col_idx, date in enumerate(forecast_dates, start=1):
        cell = hdr_cells[col_idx]
        cell.text = date.strftime('%b %d')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(6)

    # Step 5: Fill in probability data for each return period
    for row_idx, (period, row) in enumerate(rp_df.iterrows(), start=1):
        threshold = row.iloc[0]
        row_cells = table.rows[row_idx].cells

        # First column: return period label with threshold value
        cell = row_cells[0]
        cell.text = f"{period} ({threshold:,.0f})"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(7)

        # Calculate probability for each day
        for col_idx, date in enumerate(forecast_dates, start=1):
            # Get the daily maximum values for all ensemble members on this date
            daily_max_values = daily_max_df.loc[date]

            # Count how many ensemble members exceed the threshold
            num_exceeding = (daily_max_values > threshold).sum()
            num_members = len(daily_max_values)

            # Calculate probability as percentage
            probability = (num_exceeding / num_members) * 100 if num_members > 0 else 0

            # Format the cell
            cell = row_cells[col_idx]
            cell.text = f"{probability:.0f}%"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(7)
                if probability > 50:
                    run.bold = True
                    run.font.color.rgb = RGBColor(200, 0, 0)

    ## optional forecast flow values at the bottom, not sure if this is useful or not
    '''doc.add_paragraph("\n")

    # Add flow values below as readable text
    doc.add_heading('Forecast Flow Values (m³/s)', level=3)

    flow_text = ""
    for date in forecast_dates:
        try:
            target_date = pd.to_datetime(date).normalize()
            daily_flows = forecast_df[forecast_df.index.normalize() == target_date]['flow_median']

            if len(daily_flows) > 0:
                flow_value = daily_flows.mean()
                date_str = pd.to_datetime(date).strftime('%b %d')
                flow_text += f"{date_str}: {flow_value:,.0f}  |  "
        except:
            pass

    # Add as a paragraph with nice formatting
    flow_para = doc.add_paragraph(flow_text.rstrip("  |  "))
    flow_para.style = 'Normal'

    doc.add_paragraph("\n")'''


# ==========================================
# PARALLEL WORKERS
# ==========================================

def _process_forecast_task(r, d):
    """Worker to fetch forecast data and render plot to bytes."""
    formatted_date = pd.to_datetime(str(d)).strftime('%Y%m%d')
    data = forecast(river_id=r, date=formatted_date)
    fig = plots.forecast(data, plot_titles=["", f"Forecast for River: {r}"])
    img_bytes = fig.to_image(format="png", width=1000, height=600, scale=2)
    return io.BytesIO(img_bytes)


def _process_return_period_task(r, d):
    """Worker to fetch all data needed for the return period comparison."""
    formatted_date = pd.to_datetime(str(d)).strftime('%Y%m%d')
    rp_df = return_periods(river_id=r)
    forecast_data = forecast(river_id=r, date=formatted_date)
    ensemble_data = forecast_ensembles(river_id=r, date=formatted_date)

    # Generate Plot Image
    fig = plots.forecast(df=forecast_data, rp_df=None)
    img_bytes = fig.to_image(format="png", width=1000, height=600, scale=2)
    img_stream = io.BytesIO(img_bytes)

    return {
        "river_id": r, "date": d, "img_stream": img_stream,
        "rp_df": rp_df, "ensemble_data": ensemble_data,
        "forecast_data": forecast_data
    }


# ==========================================
# MAIN REPORT FUNCTIONS
# ==========================================

def forecast_report(riverids, dates, output_format='docx'):
    if not isinstance(riverids, list): riverids = [riverids]
    if not isinstance(dates, list): dates = [dates]

    print("Generating Forecast Report (Parallel Processing)...")

    # Run tasks in parallel to speed up API calls and Image Rendering
    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = []
        for d in dates:
            for r in riverids:
                futures.append(executor.submit(_process_forecast_task, r, d))

        image_streams = [f.result() for f in futures]

    report_date = dates[0] if len(dates) == 1 else f"{dates[0]}_to_{dates[-1]}"
    return _save_plots_to_file(image_streams, f"forecast_report_{report_date}.docx", "Forecast Report", output_format,
                               input_is_bytes=True)


def retrospective_report(riverids, output_format='docx'):
    if not isinstance(riverids, list): riverids = [riverids]

    # Standard sequential processing for retrospective (can be parallelized if needed similar to above)
    figures = []
    for r in riverids:
        data = retrospective(river_id=r)
        fig = plots.retrospective(data, plot_titles=["", f"Retrospective for River: {r}"])
        figures.append(fig)

    return _save_plots_to_file(figures, "retrospective_report.docx", "Retrospective Report", output_format)


def in_depth_retro(riverid, output_format='docx'):
    # Standard sequential processing
    daily = retro_daily(riverid)
    monthly = retro_monthly(riverid)
    yearly = retro_yearly(riverid)

    fig1 = plots.daily_averages(daily, plot_titles=["", f"Daily Averages for River {riverid}"])
    fig2 = plots.monthly_averages(monthly, plot_titles=["", f"Monthly Averages for River {riverid}"])
    fig3 = plots.annual_averages(yearly, plot_titles=["", f"Annual Averages for River {riverid}"])

    return _save_plots_to_file([fig1, fig2, fig3], f"in_depth_report_{riverid}.docx", "In Depth Retro", output_format)


def return_period_comparison(riverids, dates, output_format='docx'):
    if not isinstance(riverids, list): riverids = [riverids]
    if not isinstance(dates, list): dates = [dates]

    print("Generating Return Period Report (Parallel Processing)...")

    # 1. Fetch Data & Render Images in Parallel
    results = []
    with ThreadPoolExecutor(max_workers=5) as executor:
        # Submit all tasks
        futures_map = {executor.submit(_process_return_period_task, r, d): (r, d) for r in riverids for d in dates}
        for future in futures_map:
            try:
                results.append(future.result())
            except Exception as e:
                print(f"Error processing a river/date combo: {e}")

    # 2. Build Document Sequentially
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    _add_cover_page(doc, "Return Period Comparison")

    figure_num = 1
    for res in results:
        # Add Image
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(res['img_stream'], width=Inches(6.0))

        caption = doc.add_paragraph(f"Figure {figure_num}: River {res['river_id']} - Forecast for {res['date']}")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n")

        # Add Table
        _add_return_period_table(doc, res['rp_df'], res['ensemble_data'], res['forecast_data'])

        doc.add_heading(f'Comments for Figure {figure_num}', level=2)
        doc.add_paragraph("Notes:")
        doc.add_page_break()
        figure_num += 1

    doc.add_heading('Overall Comments', level=1)
    doc.add_paragraph("Please add any additional notes regarding this report below:")
    doc.add_paragraph("\n" * 10)

    # Save logic
    downloads_path = str(Path.home() / "Downloads")
    report_date = dates[0] if len(dates) == 1 else f"{dates[0]}_to_{dates[-1]}"
    filename = f"return_period_comparison_{report_date}.docx"
    if not filename.endswith('.docx'): filename += '.docx'
    docx_path = os.path.join(downloads_path, filename)

    doc.save(docx_path)
    print(f"DOCX Saved: {docx_path}")

    if output_format.lower() == 'pdf':
        print("Converting to PDF...")
        pdf_path = docx_path.replace('.docx', '.pdf')
        try:
            convert(docx_path, pdf_path)
            os.remove(docx_path)
            return os.path.abspath(pdf_path)
        except Exception as e:
            print(f"PDF conversion failed: {e}")
            return os.path.abspath(docx_path)
    else:
        return os.path.abspath(docx_path)


def fdc_curves(riverids, output_format='docx'):
    if not isinstance(riverids, list):
        riverids = [riverids]

    figures = []
    for r in riverids:
        data = fdc(river_id=r)
        fig = plots.flow_duration_curve(df=data, plot_titles=[f" for River {r}"])
        figures.append(fig)

    return _save_plots_to_file(figures, "fdc_curves_report.docx", "Flow Duration Curves", output_format)